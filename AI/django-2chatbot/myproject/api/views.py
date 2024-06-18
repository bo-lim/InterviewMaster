from django.shortcuts import render
from django.http import JsonResponse
from rest_framework import viewsets
from rest_framework.views import APIView
from rest_framework.response import Response

from .serializers import *
from api.models import *
import time
import os
from openai import OpenAI
import myproject.settings as settings
import io
from PyPDF2 import PdfReader
from docx import Document
import boto3
import re

client = OpenAI(
    api_key = settings.OPEN_API_KEY
)

assistant_id = settings.ASSISTANT_ID
chatbot_assistant_id = settings.CHATBOT_ASSISTANT_ID
s3_client = boto3.client(
    's3',
    aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
    aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
    region_name=settings.AWS_REGION
)

class coverletterAPI(APIView):
    def parsing(self, url):
        def extract_text_from_pdf(pdf_content):
            pdf_reader = PdfReader(io.BytesIO(pdf_content))
            text = ''
            for page in pdf_reader.pages:
                text += page.extract_text()
            return text
        
        def extract_text_from_docx(docx_content):
            doc = Document(io.BytesIO(docx_content))
            text = ''
            for para in doc.paragraphs:
                text += para.text
            return text
        
        def extract_text_from_txt(txt_content):
            return txt_content.decode('utf-8')
        
        def parse_s3_url(url):
            if url.startswith('s3://'):
                url = url[5:]  # "s3://" 부분 제거
                parts = url.split('/', 1) # 한 번만 분할
                bucket_name = parts[0]
                key = parts[1] if len(parts) > 1 else ''
            else:
                raise ValueError('Unsupported URL format')      
            return bucket_name, key
        
        bucket_name, key = parse_s3_url(url)
        file_obj = s3_client.get_object(Bucket=bucket_name, Key=key)
        file_content = file_obj["Body"].read().strip()

        if key.endswith('.pdf'):
            text = extract_text_from_pdf(file_content)
        elif key.endswith('.docx'):
            text = extract_text_from_docx(file_content)
        elif key.endswith('.txt'):
            text = extract_text_from_txt(file_content)
        else:
            raise ValueError('Unsupported file type')
        return text
    def post(self, request):
        coverletter_url = request.data.get('coverletter_url')
        position = request.data.get('position')
        # coverletter_url = 's3://simulation-userdata/coverletter/ygang4546@gmail.com_1718666736269_test.docx'
        # position_url= 's3://simulation-userdata/position/test.txt'

        if not coverletter_url :
            return Response({'response': 'URLs are missing'}, status=400)
        
        thread = client.beta.threads.create()
        request.session['thread_id'] = thread.id

        self_intro_text = self.parsing(coverletter_url)
        # job_desc_text = self.parsing(position) 
        
        prompt = f"자기소개서: {self_intro_text}\n직무: {position}"

        client.beta.threads.messages.create(
            thread_id=thread.id,
            role="user",
            content= prompt.replace('\n', ' ')
        )
        run = client.beta.threads.runs.create(
            thread_id=thread.id,
            assistant_id=assistant_id
        )
        run = self.wait_on_run(run, thread)

        try:
            messages = client.beta.threads.messages.list(thread_id=thread.id, order="asc")
            messages_list = list(messages)
        except:
            return Response({'response': 'Error'})
        
        if messages_list:
            assistant_response=messages_list[-1].content[0].text.value
            tts, question = self.extract_question(assistant_response)
            stop = 1 if "stop" in assistant_response else 0
            return Response({'response': tts,'question': question,'stop': stop})
        else:
            return Response({'response': 'No messages','question': 'No messages','stop': 0})
        
    def extract_question(self, response):
        # 정규 표현식으로 tts와 question을 분리
        # match = re.search(r'(질문\s?\d+:|꼬리질문:|질문:)(.*)', response)
        # if match:
        #     tts = match.group(2).strip()
        # else:
        #     tts = ""
        split_text = re.split(r'질문:|꼬리질문:|질문 1:|질문 2:|질문 3:', response)
        split_text = [text.strip("*").strip() for text in split_text]

        question = split_text[1]
        tts = split_text[0] + split_text[1]
        return tts, question
        
    def wait_on_run(self, run, thread):
        while run.status == "queued" or run.status == "in_progress":
            run = client.beta.threads.runs.retrieve(
                thread_id=thread.id,
                run_id=run.id,
            )
            time.sleep(0.5)
        return run

class chatAPI(APIView):
    def parsing(self, url):
        def extract_text_from_pdf(pdf_content):
            pdf_reader = PdfReader(io.BytesIO(pdf_content))
            text = ''
            for page in pdf_reader.pages:
                text += page.extract_text()
            return text
        
        def extract_text_from_docx(docx_content):
            doc = Document(io.BytesIO(docx_content))
            text = ''
            for para in doc.paragraphs:
                text += para.text
            return text
        
        def extract_text_from_txt(txt_content):
            return txt_content.decode('utf-8')
        
        def parse_s3_url(url):
            if url.startswith('s3://'):
                url = url[5:]  # "s3://" 부분 제거
                parts = url.split('/', 1) # 한 번만 분할
                bucket_name = parts[0]
                key = parts[1] if len(parts) > 1 else ''
            else:
                raise ValueError('Unsupported URL format')      
            return bucket_name, key
        
        bucket_name, key = parse_s3_url(url)
        file_obj = s3_client.get_object(Bucket=bucket_name, Key=key)
        file_content = file_obj["Body"].read().strip()

        if key.endswith('.pdf'):
            text = extract_text_from_pdf(file_content)
        elif key.endswith('.docx'):
            text = extract_text_from_docx(file_content)
        elif key.endswith('.txt'):
            text = extract_text_from_txt(file_content)
        else:
            raise ValueError('Unsupported file type')
        return text
    
    def post(self, request):
        text_url = request.data.get('text_url')
        # text_url = 's3://simulation-userdata/text/test.txt'

        if not text_url:
            return Response({'response': text_url}, status=400)
        
        if 'thread_id' not in request.session :
            return Response({'response': 'No thread'})
        
        thread_id = request.session['thread_id']
        thread = client.beta.threads.retrieve(thread_id=thread_id)

        text_text = self.parsing(text_url)
        
        prompt = f"대답: {text_text}"

        client.beta.threads.messages.create(
            thread_id=thread.id,
            role="user",
            content=prompt.replace('\n', ' ')
        )
        run = client.beta.threads.runs.create(
            thread_id=thread.id,
            assistant_id=assistant_id
        )
        run = self.wait_on_run(run, thread)
        try:
            messages = client.beta.threads.messages.list(thread_id=thread_id, order="asc")
            messages_list = list(messages)
        except:
            return Response({'response': 'Error'})
        
        if messages_list:
            assistant_response=messages_list[-1].content[0].text.value
            tts, question = self.extract_question(assistant_response)
            stop = 1 if "stop" in assistant_response else 0
            return Response({'response': tts,'question': question,'stop': stop})
        else:
            return Response({'response': 'No messages', 'stop': 0})
        
    def extract_question(self, response):
        # 정규 표현식으로 tts와 question을 분리
        # match = re.search(r'(질문\s?\d+:|꼬리질문:|질문:)(.*)', response)
        # if match:
        #     tts = match.group(2).strip()
        # else:
        #     tts = ""
        split_text = re.split(r'질문:|꼬리질문:|질문 1:|질문 2:|질문 3:', response)
        split_text = [text.strip("*").strip() for text in split_text]

        question = split_text[1]
        tts = split_text[0] + split_text[1]
        return tts, question
    
    def wait_on_run(self, run, thread):
        while run.status == "queued" or run.status == "in_progress":
            run = client.beta.threads.runs.retrieve(
                thread_id=thread.id,
                run_id=run.id,
            )
            time.sleep(0.5)
        return run   
        
class chatbotAPI(APIView):
    def post(self, request):
        text_content = request.data.get('text_content')

        if not text_content:
            return Response({'response': 'Text content is missing'}, status=400)
        
        if 'thread_id' not in request.session:
            return Response({'response': 'No thread'}, status=400)
        
        thread = client.beta.threads.create()
        request.session['thread_id'] = thread.id

        text = text_content
        
        prompt = f"대답: {text}"

        client.beta.threads.messages.create(
            thread_id=thread.id,
            role="user",
            content=prompt.replace('\n', ' ')
        )
        run = client.beta.threads.runs.create(
            thread_id=thread.id,
            assistant_id=chatbot_assistant_id
        )
        run = self.wait_on_run(run, thread)
        
        try:
            messages = client.beta.threads.messages.list(thread_id=thread.id, order="asc")
            messages_list = list(messages)
        except:
            return Response({'response': 'Error'})
        
        if messages_list:
            assistant_response = messages_list[-1].content[0].text.value
            stop = 1 if "stop" in assistant_response else 0
            return Response({'response': assistant_response, 'stop': stop})
        else:
            return Response({'response': 'No messages', 'stop': 0})
        
    def wait_on_run(self, run, thread):
        while run.status == "queued" or run.status == "in_progress":
            run = client.beta.threads.runs.retrieve(
                thread_id=thread.id,
                run_id=run.id,
            )
            time.sleep(0.5)
        return run